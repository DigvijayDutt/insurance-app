from flask import Flask, render_template, request, send_from_directory
import pandas as pd
from docx import Document
from docx.shared import Inches
import os
import shutil

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'

os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

@app.route('/', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        clean_folder(UPLOAD_FOLDER)
        clean_folder(OUTPUT_FOLDER)

        file = request.files['data']
        images = request.files.getlist('images')

        filepath = os.path.join(UPLOAD_FOLDER, file.filename)
        file.save(filepath)

        img_path = []

        for img in images:
            save_path  = os.path.join(UPLOAD_FOLDER, img.filename)
            img_path.append(save_path)
            os.makedirs(os.path.dirname(save_path), exist_ok=True)
            img.save(save_path)
        # Process the Excel and generate Word files
        print(img_path)
        filenames = generate_docs(filepath, img_path)
        return render_template("result.html", files=filenames)

    return render_template('index.html')

def generate_docs(excel_path, image_path):
    df = pd.read_excel(excel_path)
    columns = df.columns
    rooms = ["LIVING ROOM", "BEDROOM", "KITCHEN", "STORAGE"]

    image_paths = image_path
    img_idx = 0

    data = df.values
    filenames = []

    for j in range(len(data)):
        doc = Document("./templates/template.docx")
        doc.add_heading("FIRST INSPECTION REPORT", level=1)
        doc.add_picture("./static/home.jpg", width=Inches(2.5))

        for i in range(len(columns)):
            doc.add_heading(f"{columns[i]}", level=2)
            doc.add_paragraph(str(data[j][i]))

        doc.add_heading("PHOTOGRAPHS", level=2)
        for room in rooms:
            doc.add_heading(room, level=3)
            table = doc.add_table(rows=2, cols = 2)
            table.autofit = True
            for l in range(2):
                row_cells = table.rows[l].cells
                for m in range(2):
                    if img_idx < len(image_paths):
                        cell = row_cells[m]
                        p = cell.paragraphs[0]
                        run = p.add_run()
                        run.add_picture(image_paths[img_idx], width= Inches(2) ,height=Inches(2.5))
                        img_idx += 1

        filename = f"claim_{j}.docx"
        output_path = os.path.join(OUTPUT_FOLDER, filename)
        doc.save(output_path)
        filenames.append(filename)

    return filenames

def clean_folder(folder_path):
    for filename in os.listdir(folder_path):
        file_path = os.path.join(folder_path, filename)
        try:
            if os.path.isfile(file_path) or os.path.islink(file_path):
                os.unlink(file_path)  # remove file or symbolic link
            elif os.path.isdir(file_path):
                shutil.rmtree(file_path)  # remove sub-directory
        except Exception as e:
            print(f"Failed to delete {file_path}. Reason: {e}")

@app.route('/download/<filename>')
def download_file(filename):
    return send_from_directory(OUTPUT_FOLDER, filename, as_attachment=True)

if __name__ == '__main__':
    port = int(os.environ.get("PORT", 5000))
    app.run(host='0.0.0.0', port=port)
